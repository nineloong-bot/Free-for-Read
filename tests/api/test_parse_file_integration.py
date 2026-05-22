from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from fastapi.testclient import TestClient

from free_for_read.api.app import create_app


def build_minimal_epub() -> bytes:
    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "META-INF/container.xml",
            """<container xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
            <rootfiles><rootfile full-path="OPS/package.opf"/></rootfiles>
            </container>""",
        )
        archive.writestr(
            "OPS/package.opf",
            """<package xmlns="http://www.idpf.org/2007/opf"
            xmlns:dc="http://purl.org/dc/elements/1.1/">
            <metadata><dc:title>Test</dc:title></metadata>
            <manifest>
              <item id="c1" href="c1.xhtml" media-type="application/xhtml+xml"/>
            </manifest>
            <spine><itemref idref="c1"/></spine>
            </package>""",
        )
        archive.writestr(
            "OPS/c1.xhtml",
            "<html><body><h1>Chapter</h1><p>Paragraph text.</p></body></html>",
        )
    return buffer.getvalue()


def test_parse_file_multipart_with_pdf_succeeds(tmp_path: Path) -> None:
    """Integration: parse a real PDF through the file route using multipart."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buffer = BytesIO()
    writer.write(buffer)
    pdf_bytes = buffer.getvalue()

    client = TestClient(create_app())

    response = client.post(
        "/v1/parse/file",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["metadata"]["source_type"] == "pdf"
    assert "---" in payload["markdown"]


def test_parse_file_multipart_with_epub_returns_chapter_markdown(tmp_path: Path) -> None:
    client = TestClient(create_app(storage_root=tmp_path / "storage"))

    response = client.post(
        "/v1/parse/file",
        files={"file": ("book.epub", build_minimal_epub(), "application/epub+zip")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["metadata"]["source_type"] == "epub"
    assert "# Chapter" in payload["markdown"]
    assert "Paragraph text." in payload["markdown"]


def test_parse_file_path_mode_with_docx(tmp_path: Path) -> None:
    """Integration: parse a real DOCX through the file route using path mode."""
    import tempfile

    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Body text.")

    # Create file under cwd so the app's allowed_roots check passes
    with tempfile.TemporaryDirectory(dir=Path.cwd()) as td:
        file_path = Path(td) / "test.docx"
        doc.save(str(file_path))

        client = TestClient(create_app())

        response = client.post("/v1/parse/file", json={"path": str(file_path)})

        assert response.status_code == 200
        payload = response.json()
        assert payload["metadata"]["source_type"] == "word"
        assert payload["markdown"] == "# Title\n\nBody text."


def test_parse_and_library_routes_coexist(tmp_path: Path) -> None:
    """Integration: both /v1/parse/file and /v1/books/import work in same app."""
    storage_root = tmp_path / "storage"
    client = TestClient(create_app(storage_root=storage_root))

    # Import an EPUB through library
    import_response = client.post(
        "/v1/books/import",
        files={"file": ("book.epub", build_minimal_epub(), "application/epub+zip")},
    )
    assert import_response.status_code == 200
    book_id = import_response.json()["book"]["id"]
    chapter_id = import_response.json()["chapters"][0]["id"]

    # Verify chapter is readable
    chapter = client.get(f"/v1/books/{book_id}/chapters/{chapter_id}")
    assert chapter.status_code == 200

    # Parse a PDF through file route — both routes coexist
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buf = BytesIO()
    writer.write(buf)

    parse_response = client.post(
        "/v1/parse/file",
        files={"file": ("doc.pdf", buf.getvalue(), "application/pdf")},
    )
    assert parse_response.status_code == 200
    assert parse_response.json()["metadata"]["source_type"] == "pdf"
